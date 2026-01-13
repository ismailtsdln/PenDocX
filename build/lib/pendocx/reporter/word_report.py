from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .reporters import BaseReporter
from ..models.models import Mission
from ..core.errors import ReporterError
from ..core.logger import logger

class WordReporter(BaseReporter):
    """Generates a Word (.docx) report."""
    
    def generate(self, mission: Mission, output_path: Path) -> Path:
        try:
            doc = Document()
            
            # Title
            title = doc.add_heading(f"Penetration Test Report", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(mission.project_name)
            run.font.size = Pt(24)
            run.bold = True
            
            # Mission Meta
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(f"\n\nClient: ").bold = True
            p.add_run(f"{mission.client_name}\n")
            p.add_run(f"Author: ").bold = True
            p.add_run(f"{mission.author}\n")
            p.add_run(f"Date: ").bold = True
            p.add_run(f"{mission.start_date.strftime('%Y-%m-%d')}")
            
            doc.add_page_break()
            
            # Findings Table of Contents
            doc.add_heading("Table of Findings", level=1)
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Finding'
            hdr_cells[1].text = 'Severity'
            hdr_cells[2].text = 'CVSS'
            
            for test in mission.test_cases:
                row_cells = table.add_row().cells
                row_cells[0].text = test.title
                row_cells[1].text = test.severity
                row_cells[2].text = str(test.cvss_score) if test.cvss_score else "-"
                
            doc.add_page_break()
            
            # Detailed Findings
            doc.add_heading("Detailed Findings", level=1)
            for i, test in enumerate(mission.test_cases, 1):
                doc.add_heading(f"{i}. {test.title}", level=2)
                
                p = doc.add_paragraph()
                run = p.add_run(f"Severity: {test.severity}")
                run.bold = True
                if test.cvss_score:
                    p.add_run(f" | CVSS Score: {test.cvss_score}").bold = True
                
                if test.compliance_mapping:
                    p = doc.add_paragraph()
                    p.add_run("Compliance Mapping: ").bold = True
                    p.add_run(", ".join(test.compliance_mapping))
                
                doc.add_heading("Description", level=3)
                doc.add_paragraph(test.description)
                
                doc.add_heading("Impact", level=3)
                doc.add_paragraph(test.impact)
                
                doc.add_heading("Remediation", level=3)
                doc.add_paragraph(test.remediation)
                
                if test.artifacts:
                    doc.add_heading("Artifacts", level=3)
                    for art in test.artifacts:
                        doc.add_paragraph(f"Artifact: {art.name}")
                        if art.path.exists():
                             doc.add_picture(str(art.path), width=Inches(5))
                             
            output_file = output_path / f"{mission.project_name}_report.docx"
            doc.save(output_file)
            
            logger.info(f"Word report generated at [blue]{output_file}[/blue]")
            return output_file
        except Exception as e:
            raise ReporterError(f"Failed to generate Word report: {e}")
